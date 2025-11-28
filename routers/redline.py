# redline.py (Option B — image background + overlayed replacement text -> PDF -> DOCX via soffice)
from fastapi import APIRouter, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse, FileResponse
import tempfile, os, math, numpy as np, pandas as pd, torch, re, gc, time, shutil, subprocess, logging, uuid
from sentence_transformers import SentenceTransformer, util
import fitz, pdfplumber
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import RGBColor
from typing import List, Tuple
import html as _html
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
import io

router = APIRouter(prefix="/redline", tags=["Redlining"])
logger = logging.getLogger("redline")
logging.basicConfig(level=logging.INFO)

def sanitize_for_json(data):
    if isinstance(data, float):
        if math.isnan(data) or math.isinf(data):
            return 0.0
        return data
    elif isinstance(data, dict):
        return {k: sanitize_for_json(v) for k,v in data.items()}
    elif isinstance(data, list):
        return [sanitize_for_json(v) for v in data]
    elif isinstance(data, np.ndarray):
        return sanitize_for_json(data.tolist())
    else:
        return data

# Load playbook
PLAYBOOK_PATH = "clause_playbook.csv"
if not os.path.exists(PLAYBOOK_PATH):
    raise RuntimeError(f"Clause playbook not found at {PLAYBOOK_PATH}")

playbook = pd.read_csv(PLAYBOOK_PATH)
for col in ["standard_clause","Risk_Level","Action_Required"]:
    if col not in playbook.columns:
        playbook[col] = ""
playbook["standard_clause"] = playbook["standard_clause"].fillna("").astype(str)
playbook["Risk_Level"] = playbook["Risk_Level"].fillna("Low").astype(str)
playbook["Action_Required"] = playbook["Action_Required"].fillna("").astype(str)
playbook_clauses = playbook["standard_clause"].tolist()
logger.info(f"✅ Loaded playbook ({len(playbook_clauses)} entries)")

# Load sentence-transformers model
logger.info("🔄 Loading sentence-transformers model...")
model = SentenceTransformer("nlpaueb/legal-bert-base-uncased")
playbook_emb = model.encode(playbook_clauses, convert_to_tensor=True)
try:
    playbook_emb = torch.nan_to_num(playbook_emb, nan=0.0)
except Exception:
    pass
logger.info("✅ Playbook embeddings ready.")

# ----- Text extraction & clause splitting (as before) -----
def extract_text_from_pdf(path: str, ocr_max_pages: int = 3) -> str:
    frags = []
    try:
        with fitz.open(path) as doc:
            for pnum, page in enumerate(doc, start=1):
                try:
                    t = page.get_text("text") or ""
                    if t.strip():
                        frags.append(f"\n--- Page {pnum} ---\n" + t)
                except Exception as e:
                    logger.exception("fitz page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("PyMuPDF failed: %s", e)
    try:
        with pdfplumber.open(path) as pdf:
            for pnum, page in enumerate(pdf.pages, start=1):
                try:
                    t = page.extract_text() or ""
                    if t.strip():
                        frags.append(f"\n--- Page {pnum} (pdfplumber) ---\n" + t)
                except Exception as e:
                    logger.exception("pdfplumber page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("pdfplumber failed: %s", e)
    try:
        pages = convert_from_path(path, dpi=200, first_page=1, last_page=ocr_max_pages)
        for pnum, pil in enumerate(pages, start=1):
            try:
                t = pytesseract.image_to_string(pil, config="--oem 1 --psm 6")
                if t.strip():
                    frags.append(f"\n--- Page {pnum} (OCR) ---\n" + t)
            except Exception as e:
                logger.exception("OCR page error:")
        if any(f.strip() for f in frags):
            return "\n".join(frags).strip()
    except Exception as e:
        logger.info("OCR failed: %s", e)
    return ""

def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception as e:
        logger.exception("DOCX extract failed:")
        return ""

def split_into_clauses(text: str) -> List[str]:
    if not text or not text.strip(): return []
    text = text.replace("\r", "\n")
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'-\n', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[\.\?\!])\s+', text)
    clauses, current = [], ""
    for s in sentences:
        s = s.strip()
        if not s: continue
        if len(s.split()) < 6 and current:
            current += " " + s
        else:
            if current: clauses.append(current.strip())
            current = s
    if current: clauses.append(current.strip())
    clauses = [c for c in clauses if len(c.split()) >= 3]
    if not clauses: clauses = [text]
    return clauses

# ----- Helper: convert pdf -> docx using soffice (LibreOffice) -----
def convert_pdf_to_docx_with_soffice(pdf_path: str, out_dir: str = None, timeout: int = 60) -> str | None:
    try:
        if not out_dir:
            out_dir = os.path.dirname(pdf_path) or tempfile.gettempdir()
        os.makedirs(out_dir, exist_ok=True)
        soffice = os.environ.get("SOFFICE_PATH") or shutil.which("soffice")
        if not soffice:
            logger.warning("soffice not found in PATH. Set SOFFICE_PATH or install LibreOffice.")
            return None
        logger.info("Running soffice conversion: %s --headless --convert-to docx --outdir %s %s", soffice, out_dir, pdf_path)
        proc = subprocess.run([soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, pdf_path],
                              stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
        if proc.returncode != 0:
            logger.warning("soffice conversion exit %s; stderr=%s stdout=%s", proc.returncode, proc.stderr.decode("utf-8", errors="ignore"), proc.stdout.decode("utf-8", errors="ignore"))
            return None
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        candidate = os.path.join(out_dir, base + ".docx")
        if os.path.exists(candidate):
            logger.info("soffice conversion succeeded -> %s", candidate)
            return os.path.abspath(candidate)
        now = time.time()
        for fn in os.listdir(out_dir):
            if fn.lower().endswith(".docx"):
                full = os.path.join(out_dir, fn)
                if now - os.path.getmtime(full) < 120:
                    logger.info("soffice conversion found docx file -> %s", full)
                    return os.path.abspath(full)
        logger.warning("soffice conversion completed but could not find .docx output in %s", out_dir)
        return None
    except subprocess.TimeoutExpired:
        logger.exception("soffice conversion timed out")
        return None
    except Exception:
        logger.exception("soffice conversion failed")
        return None

# ----- KEY: Build edited PDF (image background + overlay replaced text) -----
def build_overlay_pdf_from_replacements(orig_pdf_path: str, replacements: dict, dpi:int=200) -> str:
    """
    orig_pdf_path: input PDF
    replacements: dict mapping original_clause -> suggested_clause
    Returns path to edited PDF
    """
    tmpdir = tempfile.mkdtemp(prefix="redline_overlay_")
    try:
        # 1. Convert each PDF page to an image (PIL) at given DPI
        pil_pages = convert_from_path(orig_pdf_path, dpi=dpi)
        # 2. Use pdfplumber to fetch word-level bounding boxes per page
        pages_word_boxes = []
        with pdfplumber.open(orig_pdf_path) as pdf:
            for page in pdf.pages:
                # words: list of dicts with 'text','x0','x1','top','bottom'
                words = page.extract_words(use_text_flow=True, keep_blank_chars=False)
                pages_word_boxes.append(words)

        # 3. For each page build replacement regions
        overlays = []  # per page list of tuples (bbox, replacement_text)
        for p_idx, words in enumerate(pages_word_boxes):
            page_overlays = []
            page_text = " ".join([w.get("text","") for w in words])
            # for each replacement original clause -> find if its seed appears on this page
            for orig_clause, sug_clause in replacements.items():
                orig_clause_clean = re.sub(r'\s+', ' ', orig_clause.strip())
                if not orig_clause_clean:
                    continue
                # take a short seed (first 40-120 chars) to locate on page
                seed = orig_clause_clean[:120]
                if len(seed) < 6:
                    seed = orig_clause_clean[:40]
                # naive check
                if seed.strip().lower() not in page_text.lower():
                    continue
                # Find a contiguous run of words that matches seed (case-insensitive)
                seed_words = [w.strip() for w in re.split(r'\s+', seed) if w.strip()]
                if not seed_words:
                    continue
                # Search words list for starting index where contiguous words match seed_words (case-insensitive)
                matched_idx = None
                n = len(words)
                sw_len = len(seed_words)
                for i in range(n - sw_len + 1):
                    chunk = " ".join(words[j]["text"] for j in range(i, i+sw_len))
                    if chunk.strip().lower() == " ".join(seed_words).strip().lower():
                        matched_idx = i
                        break
                if matched_idx is None:
                    # fallback naive: search for first word of seed
                    first = seed_words[0].lower()
                    for i in range(n):
                        if words[i]["text"].strip().lower() == first:
                            matched_idx = i
                            break
                if matched_idx is None:
                    continue
                # Now try to expand match to cover more words to match full orig_clause
                end_idx = matched_idx + sw_len - 1
                # try to expand forward while the concatenation is a prefix of orig_clause_clean
                while end_idx + 1 < n:
                    candidate = " ".join(words[j]["text"] for j in range(matched_idx, end_idx+2))
                    if orig_clause_clean.lower().startswith(candidate.strip().lower()):
                        end_idx += 1
                    else:
                        break
                # Compose bbox union for words matched from matched_idx..end_idx
                x0 = min(words[j]["x0"] for j in range(matched_idx, end_idx+1))
                x1 = max(words[j]["x1"] for j in range(matched_idx, end_idx+1))
                top = min(words[j]["top"] for j in range(matched_idx, end_idx+1))
                bottom = max(words[j]["bottom"] for j in range(matched_idx, end_idx+1))
                page_overlays.append(((x0, top, x1, bottom), sug_clause))
            overlays.append(page_overlays)

        # 4. Create a new PDF using reportlab: draw image background then overlay replacement text
        edited_pdf_path = os.path.join(tmpdir, f"edited_{os.path.basename(orig_pdf_path)}")
        # ensure .pdf extension
        if not edited_pdf_path.lower().endswith(".pdf"):
            edited_pdf_path += ".pdf"
        c = None
        for page_idx, pil_page in enumerate(pil_pages):
            # page size in pixels
            img_w_px, img_h_px = pil_page.size
            # convert pixels->points for reportlab: points = px * 72 / dpi
            points_w = img_w_px * 72.0 / dpi
            points_h = img_h_px * 72.0 / dpi
            if c is None:
                c = canvas.Canvas(edited_pdf_path, pagesize=(points_w, points_h))
            else:
                c.setPageSize((points_w, points_h))
            # draw page image as background
            img_stream = io.BytesIO()
            pil_page.save(img_stream, format="PNG")
            img_stream.seek(0)
            rl_image = ImageReader(img_stream)
            # draw image to fill the page (0,0) origin at bottom-left
            c.drawImage(rl_image, 0, 0, width=points_w, height=points_h)
            # overlay replacements for this page
            page_overlays = overlays[page_idx] if page_idx < len(overlays) else []
            # coordinate transform: pdfplumber origin (0,0) at top-left. reportlab origin at bottom-left.
            for bbox, replacement in page_overlays:
                x0, top, x1, bottom = bbox
                # convert pdfplumber page coordinates (assumed in points for 72dpi) to image pixels coordinate space:
                # pdfplumber gives coordinates in PDF user space (points). But since we used pdf2image with dpi, words coords scale with DPI.
                # pdfplumber's coordinates are in points (1/72in) with origin top-left. However convert_from_path used same DPI so we can map via ratio.
                # Find scale factors: pdfplumber page cropbox width vs image width
                # We'll compute mapping by getting page.width and page.height via pdfplumber for the page.
                try:
                    with pdfplumber.open(orig_pdf_path) as pdf:
                        pdfpage = pdf.pages[page_idx]
                        pdf_w = pdfpage.width
                        pdf_h = pdfpage.height
                except Exception:
                    pdf_w = points_w
                    pdf_h = points_h
                # scale x from pdf points to reportlab points (use points directly)
                # pdfplumber coordinates (x0, top) are in points (1/72), origin top-left.
                # reportlab bottom-left y = page_height_points - top_points
                # But our image->points units are points_w x points_h
                # So treat x0,x1,top,bottom as points as-is (they should be in same unit).
                # map:
                draw_x = x0 * (points_w / pdf_w) if pdf_w else x0
                draw_y_top = top * (points_h / pdf_h) if pdf_h else top
                draw_y_bottom = bottom * (points_h / pdf_h) if pdf_h else bottom
                draw_y = points_h - draw_y_top  # origin flip
                # choose a font size approximated from bbox height
                bbox_height_points = (draw_y_top - draw_y_bottom) if (draw_y_top>draw_y_bottom) else abs(draw_y_top-draw_y_bottom)
                if bbox_height_points <= 0:
                    # fallback small
                    font_size = 10
                else:
                    font_size = max(8, min(14, int(bbox_height_points*0.8)))
                # set font
                try:
                    c.setFont("Helvetica", font_size)
                except Exception:
                    pass
                # drawing text: we draw multi-line if replacement is long
                # compute max width available
                max_width = (x1 - x0) * (points_w / pdf_w) if pdf_w else (x1-x0)
                # simple wrap: break replacement into words and fill lines
                words = re.split(r'\s+', replacement.strip())
                line = ""
                lines = []
                for w in words:
                    test = (line + " " + w).strip()
                    tw = c.stringWidth(test, "Helvetica", font_size)
                    if tw <= max_width or not line:
                        line = test
                    else:
                        lines.append(line)
                        line = w
                if line:
                    lines.append(line)
                # draw each line, increment downward
                y_cursor = draw_y
                leading = font_size * 1.15
                for li in lines:
                    # ensure text baseline inside page
                    if y_cursor - leading < 0:
                        break
                    c.setFillColorRGB(0,0,0)  # black overlay text (could choose white depending on background)
                    c.drawString(draw_x, y_cursor - leading, li)
                    y_cursor -= leading
            c.showPage()
        if c:
            c.save()
        else:
            raise RuntimeError("Could not build overlay PDF")
        logger.info("Built overlay PDF: %s", edited_pdf_path)
        return edited_pdf_path
    except Exception:
        logger.exception("build_overlay_pdf_from_replacements failed")
        return None
    finally:
        # do not delete tmpdir immediately because we return path inside it
        pass

# ----- wrap fragments in HTML etc (same helpers) -----
def wrap_fragments_in_html(html_text: str, fragments: List[dict]) -> str:
    out = html_text
    for fr in fragments:
        orig = (fr.get("original") or "").strip()
        if not orig:
            continue
        seed = orig[:200]
        try:
            pattern = re.escape(seed)
            m = re.search(pattern, out)
            if not m:
                pattern2 = re.escape(seed[:80])
                m = re.search(pattern2, out, flags=re.IGNORECASE)
            if m:
                matched_text = out[m.start():m.end()]
                span = f'<span class="risky" data-clause-id="{fr.get("clause_id")}" data-risk="{fr.get("risk_level", "")}">{_html.escape(matched_text)}</span>'
                out = out[:m.start()] + span + out[m.end():]
            else:
                continue
        except Exception:
            continue
    return out

def build_simple_html_from_text_blocks(page_blocks: List[str]) -> str:
    body = "<div class='document'>\n"
    for i, p in enumerate(page_blocks):
        body += f"<div class='page' data-page='{i+1}'>\n"
        paras = [para.strip() for para in re.split(r'\n{2,}', p) if para.strip()]
        for para in paras:
            body += f"<p class='pblock'>{_html.escape(para)}</p>\n"
        body += "</div>\n"
    body += "</div>"
    full = (
        "<html><head><meta charset='utf-8'/>"
        "<style>"
        ".risky{ text-decoration: underline wavy red; text-underline-offset: 3px; cursor:pointer; } "
        ".risky[data-risk='Medium']{ text-decoration-color: orange; } "
        ".page{ padding:12px 18px; margin-bottom:20px; border-bottom:1px solid rgba(255,255,255,0.02); } "
        ".pblock{ margin:6px 0; white-space:pre-wrap; font-family: system-ui, -apple-system, 'Segoe UI', Roboto, 'Helvetica Neue', Arial; font-size:14px; color:#e6e6e6 }"
        "</style></head><body>" + body + "</body></html>"
    )
    return full

# ----- convert endpoint (mostly same as before) -----
@router.post("/convert/")
async def convert_to_html(file: UploadFile = File(None), payload: dict = Body(None)):
    source_path = None
    tmp_path = None
    try:
        if file:
            ext = os.path.splitext(file.filename or "upload")[1].lower() or ".pdf"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmp_path = tmp.name
            try:
                while True:
                    chunk = await file.read(1024 * 1024)
                    if not chunk:
                        break
                    tmp.write(chunk)
            finally:
                try: tmp.flush(); tmp.close()
                except: pass
            source_path = tmp_path
        elif payload and isinstance(payload, dict) and payload.get("path"):
            p = payload.get("path")
            if not os.path.exists(p):
                raise HTTPException(status_code=400, detail=f"Path not found: {p}")
            source_path = p
        else:
            raise HTTPException(status_code=400, detail="Provide either form-data file or JSON { 'path': '/server/path/to/file' }")

        _, ext = os.path.splitext(source_path)
        ext = ext.lower()

        page_blocks = []
        if ext == ".pdf":
            try:
                with fitz.open(source_path) as doc:
                    for pnum, page in enumerate(doc, start=1):
                        try:
                            blocks = page.get_text("blocks")
                            blocks.sort(key=lambda b: (round(b[1]), round(b[0])))
                            page_text_parts = []
                            for b in blocks:
                                txt = (b[4] or "").strip()
                                if txt:
                                    page_text_parts.append(txt)
                            page_blocks.append("\n\n".join(page_text_parts))
                        except Exception:
                            try:
                                t = page.get_text("text") or ""
                                page_blocks.append(t)
                            except Exception:
                                page_blocks.append("")
            except Exception as e:
                txt = extract_text_from_pdf(source_path, ocr_max_pages=3)
                if txt:
                    page_blocks = re.split(r'\n--- Page \d+ ---\n', txt)
                    page_blocks = [p for p in page_blocks if p and p.strip()]
                else:
                    page_blocks = [txt]
        elif ext == ".docx":
            txt = extract_text_from_docx(source_path)
            page_blocks = [txt]
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type for convert (use PDF or DOCX)")

        full_text = "\n\n".join([p for p in page_blocks if p and p.strip()])
        clauses_raw = split_into_clauses(full_text)

        clause_objs = []
        for idx, ctext in enumerate(clauses_raw):
            if not ctext or len(ctext.strip()) < 6:
                continue
            try:
                emb = model.encode(ctext, convert_to_tensor=True)
                emb = torch.nan_to_num(emb, nan=0.0)
                sims = util.pytorch_cos_sim(emb, playbook_emb)[0]
                sims = torch.nan_to_num(sims, nan=0.0)
                best_idx = int(torch.argmax(sims))
                score = float(sims[best_idx])
            except Exception:
                best_idx = 0; score = 0.0

            matched_standard = playbook_clauses[best_idx]
            matched_risk = str(playbook.loc[best_idx, "Risk_Level"]) if "Risk_Level" in playbook.columns else "Low"
            matched_action = str(playbook.loc[best_idx, "Action_Required"]) if "Action_Required" in playbook.columns else ""

            rl = matched_risk.lower()
            if rl == "high":
                suggested = matched_standard or ctext
                explanation = f"High-risk clause. Action required: {matched_action or 'Legal Review'}"
                highlight_color = "red"
            elif rl == "medium":
                suggested = matched_standard or ctext
                explanation = f"Medium-risk clause. Action recommended: {matched_action or 'Review'}"
                highlight_color = "orange"
            else:
                suggested = ctext
                explanation = "Low-risk clause. No change recommended."
                highlight_color = "green"

            clause_objs.append({
                "clause_id": f"c{idx+1}",
                "original_clause": ctext,
                "suggested_clause": suggested,
                "matched_clause": matched_standard,
                "risk_level": str(matched_risk),
                "action_required": matched_action,
                "similarity_score": score,
                "highlight_color": highlight_color,
                "explanation": explanation
            })

        fragments_to_wrap = [
            {"clause_id": o["clause_id"], "original": o["original_clause"], "risk_level": o["risk_level"]}
            for o in clause_objs if o["risk_level"].lower() in ("high", "medium")
        ]

        html_body = build_simple_html_from_text_blocks(page_blocks)
        html_with_spans = wrap_fragments_in_html(html_body, fragments_to_wrap)

        words = re.findall(r'\w+', full_text)
        total_words = len(words)
        uniq_clauses = len(set([c.strip() for c in clauses_raw]))
        dup_count = max(0, len(clauses_raw) - uniq_clauses)
        risk_score = sum(3 if r["risk_level"].lower() == "high" else (2 if r["risk_level"].lower() == "medium" else 1) for r in clause_objs)

        summary = {
            "total_clauses": len(clause_objs),
            "high_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "high"),
            "medium_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "medium"),
            "low_risk": sum(1 for x in clause_objs if x["risk_level"].lower() == "low"),
            "total_words": total_words,
            "unique_clauses": uniq_clauses,
            "duplicate_clauses": dup_count,
            "risk_score": risk_score
        }

        original_docx_token = None
        try:
            saved_uploaded_dir = "uploaded_contracts"
            os.makedirs(saved_uploaded_dir, exist_ok=True)
            unique_token = f"{uuid.uuid4().hex}{ext}"
            saved_path = os.path.join(saved_uploaded_dir, unique_token)
            try:
                shutil.copy(source_path, saved_path)
                if ext == ".pdf":
                    conv_outdir = os.path.join(tempfile.gettempdir(), f"soffice_out_{uuid.uuid4().hex}")
                    os.makedirs(conv_outdir, exist_ok=True)
                    converted = convert_pdf_to_docx_with_soffice(saved_path, out_dir=conv_outdir)
                    if converted and os.path.exists(converted):
                        conv_dest = os.path.join(saved_uploaded_dir, f"{uuid.uuid4().hex}.docx")
                        shutil.copy(converted, conv_dest)
                        original_docx_token = os.path.abspath(conv_dest)
                        logger.info("Persisted converted docx for formatting-preservation: %s", original_docx_token)
                    else:
                        # keep pdf persisted for overlay generation
                        original_docx_token = os.path.abspath(saved_path)
                        logger.warning("Could not convert uploaded PDF to DOCX for formatting-preservation; saved original PDF at %s", original_docx_token)
                else:
                    original_docx_token = os.path.abspath(saved_path)
            except Exception:
                logger.exception("Warning: could not persist source file:")
                original_docx_token = None
        except Exception:
            logger.exception("Warning while trying to persist uploaded original:")
            original_docx_token = None

        return JSONResponse(content=sanitize_for_json({
            "html": html_with_spans,
            "clauses": clause_objs,
            "summary": summary,
            "original_docx": original_docx_token
        }))

    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

# ----- generate endpoint (key: if original_docx is a PDF -> overlay approach) -----
@router.post("/generate/")
async def generate_from_clauses(payload: dict = Body(...)):
    fname = payload.get("filename") or f"edited_{int(time.time())}"
    safe_name = re.sub(r'[^\w\-_\.]', '_', str(fname))
    clauses = payload.get("clauses") or []
    original_docx = payload.get("original_docx") or None

    if not isinstance(clauses, list) or len(clauses) == 0:
        raise HTTPException(status_code=400, detail="No clauses provided to generate doc.")

    # build replacements mapping for high/medium only
    replacements = {}
    for c in clauses:
        risk = (c.get("risk_level") or "Low").lower()
        if risk in ("high", "medium"):
            orig = (c.get("original_clause") or "").strip()
            sug = (c.get("suggested_clause") or orig).strip()
            if orig:
                replacements[orig] = sug

    out_dir = "generated_reports"; os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{safe_name}.docx")

    try:
        # If original_docx provided and exists
        if original_docx and isinstance(original_docx, str) and os.path.exists(original_docx):
            try:
                _, orig_ext = os.path.splitext(original_docx)
                if orig_ext.lower() == ".pdf":
                    # Option B path: create an edited PDF overlay and convert to DOCX
                    logger.info("Original is PDF -> building overlay PDF with replacements...")
                    edited_pdf = build_overlay_pdf_from_replacements(original_docx, replacements, dpi=200)
                    if edited_pdf and os.path.exists(edited_pdf):
                        # Convert edited PDF to DOCX using soffice
                        conv_outdir = os.path.join(tempfile.gettempdir(), f"soffice_out_{uuid.uuid4().hex}")
                        os.makedirs(conv_outdir, exist_ok=True)
                        converted_docx = convert_pdf_to_docx_with_soffice(edited_pdf, out_dir=conv_outdir)
                        if converted_docx and os.path.exists(converted_docx):
                            shutil.copy(converted_docx, out_path)
                        else:
                            # fallback: create a simple docx with replaced clauses appended (last resort)
                            raise RuntimeError("Conversion of overlay PDF to DOCX failed")
                    else:
                        raise RuntimeError("Failed to build overlay PDF for replacements")
                else:
                    # original is docx: attempt structured replace (existing code)
                    doc = Document(original_docx)
                    made_change = False
                    # helper functions defined below (we expect them to be present)
                    def _replace_text_preserve_runs(paragraph, search_text, replace_text):
                        if not search_text:
                            return False
                        full = paragraph.text
                        idx = full.find(search_text)
                        if idx == -1:
                            return False
                        first_run = paragraph.runs[0] if paragraph.runs else None
                        new_text = full[:idx] + replace_text + full[idx + len(search_text):]
                        for r in paragraph.runs:
                            try:
                                r.text = ""
                            except Exception:
                                try:
                                    r._r.getparent().remove(r._r)
                                except Exception:
                                    pass
                        new_run = paragraph.add_run(new_text)
                        try:
                            if first_run:
                                new_run.bold = first_run.bold
                                new_run.italic = first_run.italic
                                try:
                                    new_run.font.size = first_run.font.size
                                except Exception:
                                    pass
                                try:
                                    new_run.font.name = first_run.font.name
                                except Exception:
                                    pass
                                try:
                                    new_run.font.underline = first_run.font.underline
                                except Exception:
                                    pass
                                try:
                                    new_run.font.color.rgb = first_run.font.color.rgb
                                except Exception:
                                    pass
                        except Exception:
                            pass
                        return True

                    def _replace_in_paragraphs(docobj, replacements):
                        changed=False
                        for para in docobj.paragraphs:
                            for orig,repl in replacements.items():
                                if orig and orig in para.text:
                                    _replace_text_preserve_runs(para, orig, repl)
                                    changed=True
                        return changed
                    def _replace_in_tables(docobj, replacements):
                        changed=False
                        for table in docobj.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        for orig,repl in replacements.items():
                                            if orig and orig in para.text:
                                                _replace_text_preserve_runs(para, orig, repl)
                                                changed=True
                        return changed
                    def _replace_in_headers_footers(docobj, replacements):
                        changed=False
                        try:
                            for section in docobj.sections:
                                header = section.header
                                for para in header.paragraphs:
                                    for orig,repl in replacements.items():
                                        if orig and orig in para.text:
                                            _replace_text_preserve_runs(para, orig, repl)
                                            changed=True
                                footer = section.footer
                                for para in footer.paragraphs:
                                    for orig,repl in replacements.items():
                                        if orig and orig in para.text:
                                            _replace_text_preserve_runs(para, orig, repl)
                                            changed=True
                        except Exception:
                            pass
                        return changed

                    p_changed = _replace_in_paragraphs(doc, replacements)
                    t_changed = _replace_in_tables(doc, replacements)
                    hf_changed = _replace_in_headers_footers(doc, replacements)
                    made_change = p_changed or t_changed or hf_changed
                    if made_change:
                        doc.save(out_path)
                    else:
                        # fallback to assembly if nothing matched
                        assembled_lines = []
                        for c in clauses:
                            orig = c.get("original_clause","") or ""
                            risk = (c.get("risk_level") or "Low").lower()
                            suggested = c.get("suggested_clause","") or orig
                            if risk in ("high","medium"):
                                assembled_lines.append(suggested.strip())
                            else:
                                assembled_lines.append(orig.strip())
                        full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])
                        new_doc = Document()
                        new_doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                        new_doc.add_paragraph("")
                        for para in full_text.split("\n\n"):
                            p = new_doc.add_paragraph()
                            p.add_run(para)
                        new_doc.save(out_path)
            except Exception:
                logger.exception("Error while trying structured replace in original_docx:")
                # fallback to plain docx assembly
                assembled_lines = []
                for c in clauses:
                    orig = c.get("original_clause","") or ""
                    risk = (c.get("risk_level") or "Low").lower()
                    suggested = c.get("suggested_clause","") or orig
                    if risk in ("high","medium"):
                        assembled_lines.append(suggested.strip())
                    else:
                        assembled_lines.append(orig.strip())
                full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])
                new_doc = Document()
                new_doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                new_doc.add_paragraph("")
                for para in full_text.split("\n\n"):
                    p = new_doc.add_paragraph()
                    p.add_run(para)
                new_doc.save(out_path)
        else:
            # No original_docx provided — fallback assembly into DOCX
            assembled_lines = []
            for c in clauses:
                orig = c.get("original_clause") or ""
                risk = (c.get("risk_level") or "Low").lower()
                suggested = c.get("suggested_clause") or orig
                if risk in ("high", "medium"):
                    assembled_lines.append(suggested.strip())
                else:
                    assembled_lines.append(orig.strip())
            full_text = "\n\n".join([ln for ln in assembled_lines if ln is not None])
            doc = Document()
            doc.add_paragraph(f"Edited document generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")
            for para in full_text.split("\n\n"):
                p = doc.add_paragraph()
                p.add_run(para)
            doc.save(out_path)
    except Exception:
        logger.exception("generate_from_clauses -> failed to build DOCX:")
        raise HTTPException(status_code=500, detail="Failed to create edited DOCX.")

    return JSONResponse(content=sanitize_for_json({
        "download_link": f"/redline/download/{os.path.basename(out_path)}"
    }))

@router.get("/download/{filename}")
def download_redlined(filename: str):
    path = os.path.join("generated_reports", filename)
    if not os.path.exists(path):
        raise HTTPException(404, "File not found")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

# ---------- END ----------
